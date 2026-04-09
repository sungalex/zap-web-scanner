"""DOCX 보고서 생성 (python-docx)"""

from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── 상수 ──

HEADING_COLORS = {"Heading 1": "1A237E", "Heading 2": "303F9F"}
VERDICT_COLORS = {"취약": "C62828", "주의": "E65100", "양호": "2E7D32", "수동점검 필요": "6A1B9A"}
VERDICT_BG_COLORS = ["FFEBEE", "FFF3E0", "E8F5E9", "F3E5F5"]
HEADER_BG = "1A237E"
KV_KEY_BG = "E8EAF6"


# ── 헬퍼 함수 ──

def set_cell_shading(cell, color_hex: str):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def verdict_color(v: str) -> str:
    return VERDICT_COLORS.get(v, "000000")


def _setup_styles(doc):
    """문서 스타일 설정"""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    for level, color in HEADING_COLORS.items():
        h = doc.styles[level]
        h.font.name = "Arial"
        h.font.size = Pt(16 if level == "Heading 1" else 12)
        h.font.bold = True
        h.font.color.rgb = RGBColor.from_string(color)


def _add_kv_table(doc, items: list):
    """키-값 테이블 (좌측 키 볼드+배경, 우측 값)"""
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (key, val) in enumerate(items):
        kc = table.rows[i].cells[0]
        kc.text = key
        kc.paragraphs[0].runs[0].bold = True
        set_cell_shading(kc, KV_KEY_BG)
        table.rows[i].cells[1].text = val or ""
    return table


def _add_header_row(table, headers: list, row_idx: int = 0):
    """테이블 헤더 행 스타일 적용"""
    for i, header in enumerate(headers):
        hc = table.rows[row_idx].cells[i]
        hc.text = header
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hc.paragraphs[0].runs[0].bold = True
        set_cell_shading(hc, HEADER_BG)
        hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ── 보고서 섹션 ──

def _add_cover_page(doc, target_url: str, date_str: str):
    """표지"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("웹 애플리케이션 취약점 분석·평가 보고서")
    run.bold = True
    run.font.size = Pt(18)

    for text in ["주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 기준",
                 "Web Application(웹) 21개 점검 항목",
                 f"점검일: {date_str}  |  대상: {target_url}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def _add_overview_section(doc, target_url: str, duration: str):
    """1. 점검 개요"""
    doc.add_heading("1. 점검 개요", level=1)
    _add_kv_table(doc, [
        ("점검 대상", target_url),
        ("점검 도구", "OWASP ZAP + Ollama Gemma 4 (로컬 AI 분석)"),
        ("점검 방법", "ZAP Spider 크롤링 → Passive/Active Scan → 수동 점검 → AI 분석"),
        ("점검 기준", "2026 주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 - Web Application(웹) 21개 항목"),
        ("소요 시간", duration),
    ])


def _add_summary_section(doc, findings: list):
    """2. 점검 결과 요약"""
    doc.add_heading("2. 점검 결과 요약", level=1)

    v_counts = {"취약": 0, "주의": 0, "양호": 0, "수동점검 필요": 0}
    for f in findings:
        v_counts[f.verdict] = v_counts.get(f.verdict, 0) + 1

    # 판정 분포 테이블
    labels = ["취약", "주의", "양호", "수동점검 필요", "합계"]
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _add_header_row(table, labels)

    values = [v_counts["취약"], v_counts["주의"], v_counts["양호"],
              v_counts["수동점검 필요"], len(findings)]
    bg_colors = VERDICT_BG_COLORS + [None]
    for i, (val, bg) in enumerate(zip(values, bg_colors)):
        vc = table.rows[1].cells[i]
        vc.text = f"{val}건"
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc.paragraphs[0].runs[0].bold = True
        if bg:
            set_cell_shading(vc, bg)

    doc.add_paragraph()

    # 전체 항목별 판정 현황
    p = doc.add_paragraph()
    run = p.add_run("전체 항목별 판정 현황:")
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(findings), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _add_header_row(table, ["코드", "점검 항목", "판정", "점검 방법"])

    for row_idx, f in enumerate(findings):
        row = table.rows[row_idx + 1]
        row.cells[0].text = f.code
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = f.name
        vc = row.cells[2]
        vc.text = f.verdict
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc.paragraphs[0].runs[0].bold = True
        vc.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(verdict_color(f.verdict))
        row.cells[3].text = f.scan_method_desc or ""

    doc.add_page_break()


def _add_detail_section(doc, findings: list):
    """3. 항목별 상세 결과"""
    doc.add_heading("3. 점검 항목별 상세 결과", level=1)
    for f in findings:
        doc.add_heading(f"{f.code} - {f.full_name}", level=2)
        rows = [("항목 코드", f.code), ("판정", f.verdict),
                ("점검 방법", f.scan_method_desc), ("상세 결과", f.detail)]
        if f.verdict in ("취약", "주의"):
            rows.append(("조치 방안", f.remediation))
        _add_kv_table(doc, rows)
        doc.add_paragraph()
    doc.add_page_break()


def _add_zap_detail_section(doc, findings: list):
    """4. ZAP 자동화 스캔 상세"""
    doc.add_heading("4. ZAP 자동화 스캔 상세 결과", level=1)

    alerts_by_type = {}
    for f in findings:
        for a in f.zap_alerts[:5]:
            key = a.get("alert", "")
            if key not in alerts_by_type:
                alerts_by_type[key] = {**a, "urls": [], "count": 0}
            alerts_by_type[key]["urls"].append(a.get("url", "")[:80])
            alerts_by_type[key]["count"] += 1

    for a in alerts_by_type.values():
        doc.add_heading(
            f"{a['alert']} ({a.get('risk', '')} - CWE-{a.get('cweid', '')})",
            level=2)
        _add_kv_table(doc, [
            ("위험도", a.get("risk", "")),
            ("CWE", f"CWE-{a.get('cweid', '')}"),
            ("탐지 건수", f"{a['count']}건"),
            ("영향 URL", ", ".join(a["urls"][:3])),
            ("설명", a.get("description", "")[:200]),
        ])
        doc.add_paragraph()
    doc.add_page_break()


def _add_conclusion_section(doc, summary: str):
    """5. 종합 의견"""
    doc.add_heading("5. 종합 의견 및 권고사항", level=1)
    for line in summary.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "※ 본 보고서는 자동화 도구(OWASP ZAP)와 로컬 AI(Ollama Gemma 4)를 활용한 "
        "진단 결과이며, 정밀 진단을 위해 수동 점검이 추가로 필요합니다.")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── 메인 ──

def generate_docx_report(findings: list, target_url: str,
                         summary: str, scan_meta: dict, output_path: str):
    """DOCX 보고서 생성"""
    doc = DocxDocument()

    # 페이지 설정 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.12)
    section.right_margin = Cm(2.12)

    _setup_styles(doc)

    date_str = datetime.now().strftime("%Y-%m-%d")
    duration = scan_meta.get("duration", "N/A")

    _add_cover_page(doc, target_url, date_str)
    _add_overview_section(doc, target_url, duration)
    _add_summary_section(doc, findings)
    _add_detail_section(doc, findings)
    _add_zap_detail_section(doc, findings)
    _add_conclusion_section(doc, summary)

    doc.save(output_path)
