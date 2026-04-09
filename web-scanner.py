#!/usr/bin/env python3
"""
주요정보통신기반시설 웹 취약점 자동 진단 시스템 v2.0
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 진단 기준: 2026 주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 - Web Application(웹) 21개 항목
- 스캔 도구: OWASP ZAP 2.17.0 (REST API)
- AI 분석:  Ollama + Gemma 4 (로컬)
- 보고서:   DOCX (가이드 양식) + HTML (대시보드) + JSON (기계판독)
- 워크플로: ZAP MCP 플레이북 8단계 기반
"""

import json
import time
import requests
import sys
import os
import re
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional, List, Dict
from urllib.parse import quote
from dotenv import load_dotenv
import logging
import warnings
from urllib3.exceptions import InsecureRequestWarning

load_dotenv()

# ============================================================
# 로깅 설정
# ============================================================
logger = logging.getLogger("web-scanner")
logger.setLevel(logging.INFO)
logging.captureWarnings(True)
warnings.filterwarnings("default")
warnings.filterwarnings("ignore", category=InsecureRequestWarning)

# ============================================================
# 설정
# ============================================================
ZAP_API_URL = os.getenv("ZAP_API_URL", "http://localhost:8090")
ZAP_API_KEY = os.getenv("ZAP_API_KEY", "")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma4:e4b")

# ============================================================
# 2026 가이드 21개 점검 항목 정의
# ============================================================
KISA_2026_ITEMS = [
    {
        "code": "CI", "name": "코드 인젝션", "full_name": "코드 인젝션 (Code Injection)",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "auto", "zap_coverage": "●",
        "description": "웹 애플리케이션 내 다양한 인젝션 공격(LDAP, 운영체제 명령 실행, SSI, XPATH, XML, SSTI 인젝션 등)에 대해 외부 입력값이 쿼리나 명령어로 삽입되어 비인가된 접근이나 코드 실행의 가능 유무 점검",
        "zap_alert_patterns": [
            "OS Command Injection", "Remote OS Command Injection",
            "LDAP Injection", "Server Side Include", "XPATH Injection",
            "XML External Entity", "XXE", "SSTI", "Template Injection",
            "Code Injection", "Remote Code Execution"
        ],
        "zap_cwe_ids": [78, 90, 97, 643, 611, 1336, 94],
    },
    {
        "code": "SI", "name": "SQL 인젝션", "full_name": "SQL 인젝션 (SQL Injection)",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "auto", "zap_coverage": "●",
        "description": "웹 애플리케이션 내 입력값이 SQL 쿼리에 삽입되어 비인가된 데이터베이스 접근과 조작 가능 여부 점검",
        "zap_alert_patterns": [
            "SQL Injection", "Blind SQL", "Time Based SQL", "Boolean Based SQL",
            "Error Based SQL", "UNION SQL"
        ],
        "zap_cwe_ids": [89],
    },
    {
        "code": "DI", "name": "디렉터리 인덱싱", "full_name": "디렉터리 인덱싱",
        "importance": "상", "category": "정보누출",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "웹 애플리케이션 서버 내 디렉터리 인덱싱 취약점 존재 여부 점검",
        "zap_alert_patterns": ["Directory Browsing", "Directory Listing"],
        "zap_cwe_ids": [548],
        "manual_check": {
            "paths": ["/_next/", "/_next/static/", "/images/", "/fonts/", "/api/", "/uploads/", "/public/"],
            "detect_patterns": ["Index of", "Directory listing", "Parent Directory"]
        }
    },
    {
        "code": "EP", "name": "에러 페이지 적용 미흡", "full_name": "에러 페이지 적용 미흡",
        "importance": "상", "category": "정보누출",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "웹 애플리케이션 에러 페이지 내 불필요한 정보 노출 여부 점검",
        "zap_alert_patterns": ["Application Error Disclosure"],
        "zap_cwe_ids": [200, 209],
        "manual_check": {
            "paths": ["/nonexistent-page-12345", "/../../etc/passwd",
                      "/<script>alert(1)</script>", "/login?id=1' OR '1'='1"],
            "detect_patterns": ["Error:", "at ", "stack", "Traceback", "node_modules"]
        }
    },
    {
        "code": "IL", "name": "정보 누출", "full_name": "정보 누출",
        "importance": "상", "category": "정보누출",
        "scan_method": "mixed", "zap_coverage": "◐",
        "description": "웹 애플리케이션 내 중요 정보 및 불필요한 정보의 노출 여부 점검",
        "zap_alert_patterns": [
            "Information Disclosure", "Private IP Disclosure",
            "Server Leaks Information", "X-Content-Type-Options",
            "Tech Detected", "Timestamp Disclosure"
        ],
        "zap_cwe_ids": [200],
    },
    {
        "code": "XS", "name": "크로스사이트 스크립팅", "full_name": "크로스사이트 스크립트",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "mixed", "zap_coverage": "◐",
        "description": "웹 애플리케이션 내 악성 스크립트가 다른 사용자의 브라우저에서 실행되는 취약점 존재 여부 점검",
        "zap_alert_patterns": [
            "Cross Site Scripting", "XSS", "Reflected XSS", "Stored XSS", "DOM XSS",
            "CSP: script-src unsafe-inline", "CSP: script-src unsafe-eval",
            "CSP: style-src unsafe-inline"
        ],
        "zap_cwe_ids": [79, 693],
    },
    {
        "code": "CF", "name": "크로스사이트 요청 위조", "full_name": "크로스사이트 요청 위조(CSRF)",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "웹 애플리케이션 내 사용자의 인증 세션을 악용하여 의도하지 않은 위조 요청 가능 여부 점검",
        "zap_alert_patterns": [
            "Cross-Site Request Forgery", "CSRF", "Anti-CSRF",
            "X-Frame-Options"
        ],
        "zap_cwe_ids": [352],
    },
    {
        "code": "SF", "name": "서버사이드 요청 위조", "full_name": "서버사이드 요청 위조(SSRF)",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "auto", "zap_coverage": "●",
        "description": "입력값을 통해 외부에서 직접적인 접근이 제한된 내부 서버 자원에 접근하여 악의적인 요청을 처리하거나 중요 정보의 유출 여부 점검",
        "zap_alert_patterns": ["Server Side Request Forgery", "SSRF"],
        "zap_cwe_ids": [918],
    },
    {
        "code": "BF", "name": "약한 비밀번호 정책", "full_name": "약한 비밀번호 정책",
        "importance": "상", "category": "인증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "회원가입 시 비밀번호 복잡도 정책의 적절성 점검",
        "zap_alert_patterns": ["Weak Authentication", "Weak Password"],
        "zap_cwe_ids": [521],
    },
    {
        "code": "IA", "name": "불충분한 인증 절차", "full_name": "불충분한 인증 절차",
        "importance": "상", "category": "인증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "인증 절차의 적절성(CAPTCHA, MFA 등) 점검",
        "zap_alert_patterns": ["Authentication"],
        "zap_cwe_ids": [287],
    },
    {
        "code": "IN", "name": "불충분한 권한 검증", "full_name": "불충분한 권한 검증",
        "importance": "상", "category": "인가",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "비인증 상태와 인증 상태의 접근 차이 및 수평/수직 권한 검증 점검",
        "zap_alert_patterns": ["Access Control", "Authorization", "IDOR"],
        "zap_cwe_ids": [285],
    },
    {
        "code": "PR", "name": "취약한 비밀번호 복구", "full_name": "취약한 비밀번호 복구 절차",
        "importance": "상", "category": "인증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "비밀번호 복구 프로세스의 본인 인증 강도 점검",
        "zap_alert_patterns": ["Password Recovery"],
        "zap_cwe_ids": [640],
    },
    {
        "code": "PV", "name": "프로세스 검증 누락", "full_name": "프로세스 검증 누락",
        "importance": "상", "category": "프로세스검증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "다단계 프로세스의 순서 우회 가능 여부 점검",
        "zap_alert_patterns": ["Process Validation"],
        "zap_cwe_ids": [841],
    },
    {
        "code": "FU", "name": "악성 파일 업로드", "full_name": "악성 파일 업로드",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "파일 업로드 기능에서 악성 파일 업로드 가능 여부 점검",
        "zap_alert_patterns": ["File Upload"],
        "zap_cwe_ids": [434],
    },
    {
        "code": "FD", "name": "파일 다운로드", "full_name": "파일 다운로드",
        "importance": "상", "category": "입력값 검증",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "파일 다운로드 경로 조작을 통한 비인가 파일 접근 가능 여부 점검",
        "zap_alert_patterns": ["Path Traversal", "Directory Traversal"],
        "zap_cwe_ids": [22],
    },
    {
        "code": "IS", "name": "불충분한 세션 관리", "full_name": "불충분한 세션 관리",
        "importance": "상", "category": "세션관리",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "세션 토큰 저장 방식(쿠키/localStorage), HttpOnly/Secure 플래그 적용 여부 점검",
        "zap_alert_patterns": [
            "Session Management", "Session Fixation", "Session ID",
            "Cookie Without HttpOnly", "Cookie Without Secure"
        ],
        "zap_cwe_ids": [384, 613, 614],
    },
    {
        "code": "SN", "name": "데이터 평문 전송", "full_name": "데이터 평문 전송",
        "importance": "상", "category": "암호화",
        "scan_method": "mixed", "zap_coverage": "◐",
        "description": "HTTPS 적용 여부, HSTS 설정 여부 등 전송 구간 보안 점검",
        "zap_alert_patterns": [
            "Strict-Transport-Security", "HSTS", "Insecure Transport",
            "SSL", "TLS", "Mixed Content"
        ],
        "zap_cwe_ids": [319, 523],
    },
    {
        "code": "CC", "name": "쿠키 변조", "full_name": "쿠키 변조",
        "importance": "상", "category": "세션관리",
        "scan_method": "mixed", "zap_coverage": "◐",
        "description": "쿠키 보안 속성(HttpOnly, Secure, SameSite) 및 CORS 설정 점검",
        "zap_alert_patterns": [
            "Cookie", "HttpOnly", "Secure Flag", "SameSite",
            "Cross-Domain Misconfiguration", "CORS"
        ],
        "zap_cwe_ids": [565, 264, 942],
    },
    {
        "code": "AE", "name": "관리자 페이지 노출", "full_name": "관리자 페이지 노출",
        "importance": "상", "category": "정보누출",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "관리자 페이지 접근 가능 여부 점검",
        "zap_alert_patterns": ["Admin Page", "Hidden File"],
        "zap_cwe_ids": [200],
        "manual_check": {
            "paths": ["/admin", "/administrator", "/admin/login", "/manage",
                      "/manager", "/cms", "/wp-admin", "/dashboard",
                      "/console", "/backoffice", "/phpmyadmin"],
        }
    },
    {
        "code": "AU", "name": "자동화 공격", "full_name": "자동화 공격",
        "importance": "상", "category": "자동화공격",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "CAPTCHA 적용 여부, 로그인 시도 제한 등 자동화 공격 방어 점검",
        "zap_alert_patterns": ["Brute Force", "CAPTCHA"],
        "zap_cwe_ids": [799, 307],
    },
    {
        "code": "WM", "name": "불필요한 Method 악용", "full_name": "불필요한 Method 악용",
        "importance": "상", "category": "설정",
        "scan_method": "manual", "zap_coverage": "○",
        "description": "불필요한 HTTP Method(OPTIONS, PUT, DELETE, PATCH) 허용 여부 점검",
        "zap_alert_patterns": ["HTTP Method"],
        "zap_cwe_ids": [749],
        "manual_check": {
            "methods": ["OPTIONS", "PUT", "DELETE", "PATCH", "HEAD"],
        }
    },
]

# ZAP Alert → 가이드 항목 매핑 규칙 (플레이북 기반)
ZAP_ALERT_MAPPING_RULES = {
    "SQL Injection": {"code": "SI", "verdict": "취약"},
    "Cross Site Scripting": {"code": "XS", "verdict": "취약"},
    "Reflected XSS": {"code": "XS", "verdict": "취약"},
    "Stored XSS": {"code": "XS", "verdict": "취약"},
    "OS Command Injection": {"code": "CI", "verdict": "취약"},
    "Code Injection": {"code": "CI", "verdict": "취약"},
    "LDAP Injection": {"code": "CI", "verdict": "취약"},
    "Server Side Request Forgery": {"code": "SF", "verdict": "취약"},
    "CSP: script-src unsafe-inline": {"code": "XS", "verdict": "주의"},
    "CSP: script-src unsafe-eval": {"code": "XS", "verdict": "주의"},
    "CSP: style-src unsafe-inline": {"code": "XS", "verdict": "주의"},
    "Cross-Domain Misconfiguration": {"code": "CC", "verdict": "주의"},
    "Strict-Transport-Security Header Not Set": {"code": "SN", "verdict": "주의"},
    "X-Content-Type-Options Header Missing": {"code": "IL", "verdict": "주의"},
    "X-Frame-Options Header Not Set": {"code": "CF", "verdict": "주의"},
    "Directory Browsing": {"code": "DI", "verdict": "취약"},
    "Application Error Disclosure": {"code": "EP", "verdict": "취약"},
    "Information Disclosure": {"code": "IL", "verdict": "주의"},
    "Cookie Without Secure Flag": {"code": "CC", "verdict": "주의"},
    "Cookie Without HttpOnly Flag": {"code": "IS", "verdict": "주의"},
    "Cookie No HttpOnly Flag": {"code": "IS", "verdict": "주의"},
}


# ============================================================
# 데이터 클래스
# ============================================================
@dataclass
class Finding:
    code: str
    name: str
    full_name: str
    importance: str
    verdict: str           # 취약, 주의, 양호, 수동점검 필요
    scan_method_desc: str  # 실제 수행한 점검 방법 설명
    detail: str            # 상세 결과
    remediation: str = ""  # 조치 방안 (취약/주의 시)
    zap_alerts: list = field(default_factory=list)
    ai_analysis: str = ""


# ============================================================
# ZAP Scanner (REST API)
# ============================================================
class ZAPScanner:
    def __init__(self, base_url=ZAP_API_URL, api_key=ZAP_API_KEY):
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key

    def _get(self, path, params=None):
        params = params or {}
        if self.api_key:
            params["apikey"] = self.api_key
        try:
            r = requests.get(f"{self.base_url}{path}", params=params, timeout=30)
            r.raise_for_status()
            return r.json()
        except Exception as e:
            logger.error("[ZAP 오류] %s: %s", path, e)
            print(f"  [ZAP 오류] {path}: {e}")
            return None

    def check(self) -> bool:
        result = self._get("/JSON/core/view/version/")
        if result:
            logger.info("[ZAP] 연결 성공 - v%s", result.get('version','?'))
            print(f"  [ZAP] 연결 성공 - v{result.get('version','?')}")
            return True
        logger.error("[ZAP] 연결 실패")
        return False

    # [Step 3] Context 생성
    def create_context(self, name: str) -> str:
        r = self._get("/JSON/context/action/newContext/", {"contextName": name})
        cid = r.get("contextId", "") if r else ""
        logger.info("[Context] 생성: %s (ID: %s)", name, cid)
        print(f"  [Context] 생성: {name} (ID: {cid})")
        return cid

    def include_in_context(self, name: str, regex: str):
        self._get("/JSON/context/action/includeInContext/",
                  {"contextName": name, "regex": regex})

    # [Step 3] 인증 설정
    def setup_auth(self, context_id: str, login_url: str, login_data: str,
                   auth_type: str = "jsonBasedAuthentication"):
        encoded_url = quote(login_url, safe='')
        encoded_data = quote(login_data, safe='')
        params = f"loginUrl={encoded_url}&loginRequestData={encoded_data}"
        self._get(f"/JSON/authentication/action/setAuthenticationMethod/",
                  {"contextId": context_id, "authMethodName": auth_type,
                   "authMethodConfigParams": params})

    def set_indicators(self, context_id: str, logged_in: str, logged_out: str):
        self._get("/JSON/authentication/action/setLoggedInIndicator/",
                  {"contextId": context_id, "loggedInIndicatorRegex": logged_in})
        self._get("/JSON/authentication/action/setLoggedOutIndicator/",
                  {"contextId": context_id, "loggedOutIndicatorRegex": logged_out})

    def create_user(self, context_id: str, username: str, password: str) -> str:
        r = self._get("/JSON/users/action/newUser/",
                      {"contextId": context_id, "name": "tester"})
        uid = r.get("userId", "") if r else ""
        if uid:
            encoded_user = quote(username, safe='')
            encoded_pass = quote(password, safe='')
            self._get("/JSON/users/action/setAuthenticationCredentials/",
                      {"contextId": context_id, "userId": uid,
                       "authCredentialsConfigParams":
                           f"username={encoded_user}&password={encoded_pass}"})
            self._get("/JSON/users/action/setUserEnabled/",
                      {"contextId": context_id, "userId": uid, "enabled": "true"})
            self._get("/JSON/forcedUser/action/setForcedUser/",
                      {"contextId": context_id, "userId": uid})
            self._get("/JSON/forcedUser/action/setForcedUserModeEnabled/",
                      {"enabled": "true"})
        return uid

    # [Step 4] Spider
    def spider(self, url: str, context_name: str = "") -> list:
        logger.info("[Spider] 크롤링 시작: %s", url)
        print(f"\n  [Spider] 크롤링 시작: {url}")
        params = {"url": url, "recurse": "true", "subtreeOnly": "true"}
        if context_name:
            params["contextName"] = context_name
        r = self._get("/JSON/spider/action/scan/", params)
        if not r:
            return []
        sid = r.get("scan", "0")
        while True:
            s = self._get("/JSON/spider/view/status/", {"scanId": sid})
            if not s:
                break
            p = int(s.get("status", "100"))
            print(f"\r  [Spider] {p}%", end="", flush=True)
            if p >= 100:
                break
            time.sleep(2)
        print()
        urls_r = self._get("/JSON/spider/view/results/", {"scanId": sid})
        urls = urls_r.get("results", []) if urls_r else []
        logger.info("[Spider] 발견 URL: %d개", len(urls))
        print(f"  [Spider] 발견 URL: {len(urls)}개")
        return urls

    # Ajax Spider (SPA용)
    def ajax_spider(self, url: str, context_name: str = ""):
        logger.info("[Ajax Spider] 시작: %s", url)
        print(f"  [Ajax Spider] 시작: {url}")
        params = {"url": url, "subtreeOnly": "true"}
        if context_name:
            params["contextName"] = context_name
        self._get("/JSON/ajaxSpider/action/scan/", params)
        while True:
            s = self._get("/JSON/ajaxSpider/view/status/")
            if not s or s.get("status", "stopped") == "stopped":
                break
            time.sleep(3)
        logger.info("[Ajax Spider] 완료")
        print("  [Ajax Spider] 완료")

    # [Step 5] Passive / Active Scan
    def wait_passive(self, timeout=120, max_retries=3):
        for attempt in range(1, max_retries + 1):
            start = time.time()
            remaining = -1
            while time.time() - start < timeout:
                r = self._get("/JSON/pscan/view/recordsToScan/")
                remaining = int(r.get("recordsToScan", "0")) if r else -1
                if remaining == 0:
                    logger.info("[Passive Scan] 완료")
                    print(f"  [Passive Scan] 완료")
                    return True
                time.sleep(3)
            # timeout 발생
            if attempt < max_retries:
                logger.warning("[Passive Scan] 타임아웃 (잔여 %s건) - 재시도 %d/%d",
                               remaining, attempt, max_retries)
                print(f"  [Passive Scan] 타임아웃 (잔여 {remaining}건) - 재시도 {attempt}/{max_retries}")
            else:
                logger.warning("[Passive Scan] 타임아웃 (잔여 %s건) - 최대 재시도 초과, 계속 진행",
                               remaining)
                print(f"  [Passive Scan] 타임아웃 (잔여 {remaining}건) - 최대 재시도 초과, 계속 진행")
        return False

    def active_scan(self, url: str, context_id: str = "") -> str:
        logger.info("[Active Scan] 시작: %s", url)
        print(f"\n  [Active Scan] 시작: {url}")
        params = {"url": url, "recurse": "true", "subtreeOnly": "true"}
        if context_id:
            params["contextId"] = context_id
        r = self._get("/JSON/ascan/action/scan/", params)
        if not r:
            return ""
        sid = r.get("scan", "0")
        while True:
            s = self._get("/JSON/ascan/view/status/", {"scanId": sid})
            if not s:
                break
            p = int(s.get("status", "100"))
            print(f"\r  [Active Scan] {p}%", end="", flush=True)
            if p >= 100:
                break
            time.sleep(5)
        print()
        logger.info("[Active Scan] 완료 (scanId: %s)", sid)
        return sid

    # [Step 6] 결과 수집
    def get_alerts(self, base_url: str = "", risk_id: str = "") -> list:
        params = {"start": "0", "count": "500"}
        if base_url:
            params["baseurl"] = base_url
        if risk_id:
            params["riskId"] = risk_id
        r = self._get("/JSON/alert/view/alerts/", params)
        return r.get("alerts", []) if r else []

    def get_alerts_summary(self, base_url: str) -> dict:
        r = self._get("/JSON/alert/view/alertsSummary/", {"baseurl": base_url})
        return r.get("alertsSummary", {}) if r else {}

    # [Step 7] 수동 점검 시뮬레이션 (HTTP 요청으로 수행)
    def manual_check_directories(self, base_url: str) -> list:
        """DI 점검: 디렉터리 인덱싱"""
        results = []
        paths = ["/_next/", "/_next/static/", "/images/", "/fonts/",
                 "/api/", "/uploads/", "/public/"]
        for path in paths:
            try:
                r = requests.get(f"{base_url.rstrip('/')}{path}", timeout=10,
                                allow_redirects=True, verify=False)
                listing = any(p in r.text for p in
                             ["Index of", "Directory listing", "Parent Directory"])
                results.append({"path": path, "status": r.status_code, "listing": listing})
            except Exception as e:
                logger.error("[DI] %s 점검 실패: %s", path, e)
                results.append({"path": path, "status": "ERR", "listing": False})
        return results

    def manual_check_error_pages(self, base_url: str) -> list:
        """EP 점검: 에러 페이지"""
        results = []
        paths = ["/nonexistent-page-12345", "/../../etc/passwd",
                 "/%3Cscript%3Ealert(1)%3C/script%3E"]
        for path in paths:
            try:
                r = requests.get(f"{base_url.rstrip('/')}{path}", timeout=10,
                                allow_redirects=True, verify=False)
                results.append({
                    "path": path, "status": r.status_code,
                    "stack_trace": any(p in r.text for p in
                                      ["Traceback", "Exception in", "Stack Trace",
                                       "at java.", "at org.", "at com.", "at net.",
                                       "in /var/", "in /home/", "in C:\\",
                                       "line \\d+", "Fatal error"]),
                    "server_info": any(p in r.text for p in ["nginx", "Apache", "Express", "node_modules"]),
                    "db_info": any(p in r.text for p in ["SQL", "mysql", "postgres"]),
                })
            except Exception as e:
                logger.error("[EP] %s 점검 실패: %s", path, e)
                results.append({"path": path, "status": "ERR"})
        return results

    def manual_check_admin_pages(self, base_url: str) -> list:
        """AE 점검: 관리자 페이지 노출"""
        results = []
        paths = ["/admin", "/administrator", "/admin/login", "/manage",
                 "/manager", "/cms", "/wp-admin", "/dashboard",
                 "/console", "/backoffice", "/phpmyadmin"]
        for path in paths:
            try:
                r = requests.get(f"{base_url.rstrip('/')}{path}", timeout=10,
                                allow_redirects=False, verify=False)
                results.append({"path": path, "status": r.status_code})
            except Exception as e:
                logger.error("[AE] %s 점검 실패: %s", path, e)
                results.append({"path": path, "status": "ERR"})
        return results

    def manual_check_http_methods(self, base_url: str) -> list:
        """WM 점검: 불필요한 HTTP Method"""
        results = []
        for method in ["OPTIONS", "PUT", "DELETE", "PATCH", "HEAD", "TRACE"]:
            try:
                r = requests.request(method, f"{base_url.rstrip('/')}/",
                                    timeout=10, verify=False)
                results.append({"method": method, "status": r.status_code})
            except Exception as e:
                logger.error("[WM] %s 점검 실패: %s", method, e)
                results.append({"method": method, "status": "ERR"})
        return results

    def manual_check_security_headers(self, base_url: str) -> dict:
        """SN 점검: 보안 헤더"""
        try:
            r = requests.get(f"{base_url.rstrip('/')}/", timeout=10, verify=False)
            headers_to_check = [
                "strict-transport-security", "content-security-policy",
                "x-content-type-options", "x-frame-options",
                "referrer-policy", "permissions-policy"
            ]
            result = {"https": base_url.startswith("https"), "headers": {}}
            for h in headers_to_check:
                result["headers"][h] = r.headers.get(h, None)
            return result
        except Exception as e:
            logger.error("[SN] 보안 헤더 점검 실패: %s", e)
            return {"https": False, "headers": {}}

    def manual_check_cookies(self, base_url: str) -> dict:
        """CC 점검: 쿠키 보안 속성"""
        try:
            r = requests.get(f"{base_url.rstrip('/')}/", timeout=10, verify=False)
            cookies = []
            for cookie_header in r.headers.get("set-cookie", "").split(","):
                if not cookie_header.strip():
                    continue
                lower = cookie_header.lower()
                cookies.append({
                    "raw": cookie_header.strip()[:200],
                    "httponly": "httponly" in lower,
                    "secure": "secure" in lower,
                    "samesite": "samesite" in lower,
                })
            return {
                "https": base_url.startswith("https"),
                "cookie_count": len(cookies),
                "cookies": cookies,
            }
        except Exception as e:
            logger.error("[CC] 쿠키 점검 실패: %s", e)
            return {"https": False, "cookie_count": 0, "cookies": []}


# ============================================================
# Gemma 4 AI Analyzer
# ============================================================
class GemmaAnalyzer:
    def __init__(self, base_url=OLLAMA_URL, model=OLLAMA_MODEL):
        self.base_url = base_url.rstrip("/")
        self.model = model

    def check(self) -> bool:
        try:
            r = requests.get(f"{self.base_url}/api/tags", timeout=10)
            models = [m["name"] for m in r.json().get("models", [])]
            if any("gemma4" in m for m in models):
                logger.info("[Ollama] 연결 성공 - %s", self.model)
                print(f"  [Ollama] 연결 성공 - {self.model}")
                return True
            logger.error("[Ollama] gemma4 미설치")
            print(f"  [Ollama] gemma4 미설치. 'ollama pull {self.model}' 실행 필요")
            return False
        except Exception as e:
            logger.error("[Ollama] 연결 실패: %s", e)
            print(f"  [Ollama] 연결 실패: {e}")
            return False

    def _chat(self, system: str, user: str, temp=0.3, show_progress=False) -> str:
        try:
            r = requests.post(f"{self.base_url}/api/chat", json={
                "model": self.model,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user}
                ],
                "stream": True,
                "options": {"temperature": temp, "num_ctx": 8192}
            }, timeout=900, stream=True)
            r.raise_for_status()
            content = ""
            token_count = 0
            start = time.time()
            for line in r.iter_lines():
                if not line:
                    continue
                chunk = json.loads(line)
                token = chunk.get("message", {}).get("content", "")
                content += token
                token_count += 1
                if show_progress and token_count % 20 == 0:
                    elapsed = int(time.time() - start)
                    print(f"\r    ⏳ 생성 중... {token_count} 토큰 ({elapsed}초 경과)", end="", flush=True)
                if chunk.get("done"):
                    break
            if show_progress and token_count > 0:
                elapsed = int(time.time() - start)
                print(f"\r    ✅ 완료: {token_count} 토큰 생성 ({elapsed}초 소요)        ")
            return content
        except Exception as e:
            logger.error("AI 분석 실패: %s", e)
            return f"AI 분석 실패: {e}"

    def analyze_item(self, item: dict, zap_alerts: list,
                     manual_results: dict = None) -> dict:
        """21개 항목별 AI 분석"""
        system = """당신은 주요정보통신기반시설 취약점 진단 전문가입니다.
2026 주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 - Web Application(웹) 21개 항목 기준으로 판정합니다.

판정 기준:
- 양호: ZAP 미탐지 + 수동 점검에서 문제 없음
- 취약: ZAP 또는 수동 점검에서 명확한 취약점 확인
- 주의: 직접적 취약점은 아니나 보안 강화 필요 (CSP 미흡, CAPTCHA 미적용 등)
- 수동점검 필요: 자동화 도구로 확인 불가, 별도 수동 테스트 필요

반드시 아래 JSON 형식으로만 응답하세요."""

        alert_text = ""
        for a in zap_alerts[:8]:
            alert_text += f"- {a.get('alert','')}: {a.get('url','')[:60]} (Risk:{a.get('risk','')}, CWE:{a.get('cweid','')})\n"

        manual_text = ""
        if manual_results:
            manual_text = f"\n수동 점검 결과:\n{json.dumps(manual_results, ensure_ascii=False, indent=1)[:500]}"

        user = f"""점검항목: [{item['code']}] {item['full_name']} (중요도: {item['importance']})
설명: {item['description']}
ZAP 커버리지: {item['zap_coverage']}
ZAP 탐지 경고 ({len(zap_alerts)}건):
{alert_text if alert_text else '관련 경고 없음'}
{manual_text}

JSON 응답:
{{"verdict":"취약/주의/양호/수동점검 필요","scan_method_desc":"실제 수행한 점검 방법","detail":"상세 결과 (3-5문장)","remediation":"조치 방안 (취약/주의 시)"}}"""

        response = self._chat(system, user)
        try:
            cleaned = response.strip()
            if "```" in cleaned:
                cleaned = cleaned.split("```")[1] if "```json" not in cleaned else cleaned.split("```json")[1]
                cleaned = cleaned.split("```")[0]
            start = cleaned.find("{")
            end = cleaned.rfind("}") + 1
            if start >= 0 and end > start:
                return json.loads(cleaned[start:end])
        except (json.JSONDecodeError, ValueError, IndexError) as e:
            logger.warning("[AI] JSON 파싱 실패 [%s]: %s", item.get("code", "?"), e)
        return {"verdict": "수동점검 필요", "scan_method_desc": "AI 분석",
                "detail": response[:400], "remediation": "수동 점검 필요"}

    def generate_summary(self, findings: list, target_url: str) -> str:
        """종합 의견"""
        vuln = [f for f in findings if f.verdict == "취약"]
        warn = [f for f in findings if f.verdict == "주의"]
        safe = [f for f in findings if f.verdict == "양호"]
        manual = [f for f in findings if f.verdict == "수동점검 필요"]

        system = "주요정보통신기반시설 웹 취약점 진단 결과 종합 의견을 작성하는 보안 컨설턴트입니다. 한국어로 작성합니다."
        user = f"""대상: {target_url}
진단일: {datetime.now().strftime('%Y-%m-%d')}
기준: 2026 주요정보통신기반시설 기술적 취약점 분석·평가 가이드 (웹 21개 항목)

판정 분포: 취약 {len(vuln)}건, 주의 {len(warn)}건, 양호 {len(safe)}건, 수동점검 필요 {len(manual)}건

취약 항목: {', '.join(f'[{f.code}]{f.name}' for f in vuln) or '없음'}
주의 항목: {', '.join(f'[{f.code}]{f.name}' for f in warn) or '없음'}

다음 구조로 종합 의견을 작성하세요:
1. 주요 발견사항 (취약/주의 항목 설명)
2. 우선순위별 조치 권고 (긴급/단기/중장기)
3. 추가 점검 권고 (수동점검 필요 항목)"""

        return self._chat(system, user, temp=0.2, show_progress=True)


# Alert 매핑 엔진
# ============================================================
def map_alerts_to_items(alerts: list) -> dict:
    """ZAP 경고를 21개 항목에 매핑"""
    mapping = {item["code"]: [] for item in KISA_2026_ITEMS}

    for alert in alerts:
        alert_name = alert.get("alert", "")
        cwe_id = int(alert.get("cweid", "0") or "0")
        matched = False

        # 1차: 플레이북 매핑 규칙 (정확 매칭)
        for pattern, rule in ZAP_ALERT_MAPPING_RULES.items():
            if pattern.lower() in alert_name.lower():
                mapping[rule["code"]].append(alert)
                matched = True
                break

        if matched:
            continue

        # 2차: CWE ID 매칭
        for item in KISA_2026_ITEMS:
            if cwe_id in item.get("zap_cwe_ids", []):
                mapping[item["code"]].append(alert)
                matched = True
                break

        if matched:
            continue

        # 3차: 키워드 매칭
        for item in KISA_2026_ITEMS:
            for pattern in item.get("zap_alert_patterns", []):
                if pattern.lower() in alert_name.lower():
                    mapping[item["code"]].append(alert)
                    matched = True
                    break
            if matched:
                break

    return mapping


# ============================================================
# 보고서 생성 (JSON)
# ============================================================
def generate_json_report(findings: list, target_url: str,
                         summary: str, scan_meta: dict) -> dict:
    return {
        "report_title": "웹 애플리케이션 취약점 분석·평가 보고서",
        "standard": "2026 주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 - Web Application(웹) 21개 항목",
        "target": target_url,
        "scan_date": datetime.now().isoformat(),
        "scan_tools": f"OWASP ZAP + Ollama {OLLAMA_MODEL}",
        "duration": scan_meta.get("duration", ""),
        "verdict_summary": {
            "취약": sum(1 for f in findings if f.verdict == "취약"),
            "주의": sum(1 for f in findings if f.verdict == "주의"),
            "양호": sum(1 for f in findings if f.verdict == "양호"),
            "수동점검 필요": sum(1 for f in findings if f.verdict == "수동점검 필요"),
            "합계": len(findings),
        },
        "findings": [
            {
                "code": f.code, "name": f.name, "full_name": f.full_name,
                "importance": f.importance, "verdict": f.verdict,
                "scan_method": f.scan_method_desc, "detail": f.detail,
                "remediation": f.remediation, "alert_count": len(f.zap_alerts),
            }
            for f in findings
        ],
        "summary": summary,
    }


# ============================================================
# DOCX 보고서 생성 (python-docx)
# ============================================================
def generate_docx_report(findings: list, target_url: str,
                         summary: str, scan_meta: dict, output_path: str):
    """python-docx를 사용하여 DOCX 보고서 생성"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = DocxDocument()

    # 페이지 설정 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.12)
    section.right_margin = Cm(2.12)

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    for level, size, color in [("Heading 1", 16, "1A237E"), ("Heading 2", 12, "303F9F")]:
        h = doc.styles[level]
        h.font.name = "Arial"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor.from_string(color)

    def set_cell_shading(cell, color_hex):
        """셀 배경색 설정"""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_kv_table(items):
        """키-값 테이블 생성 (좌측 키 볼드 + 배경, 우측 값)"""
        table = doc.add_table(rows=len(items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, (key, val) in enumerate(items):
            kc = table.rows[i].cells[0]
            kc.text = key
            kc.paragraphs[0].runs[0].bold = True
            set_cell_shading(kc, "E8EAF6")
            table.rows[i].cells[1].text = val or ""
        return table

    def verdict_color(v):
        return {"취약": "C62828", "주의": "E65100", "양호": "2E7D32",
                "수동점검 필요": "6A1B9A"}.get(v, "000000")

    date_str = datetime.now().strftime("%Y-%m-%d")
    duration = scan_meta.get("duration", "N/A")

    # === 표지 ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("웹 애플리케이션 취약점 분석·평가 보고서")
    run.bold = True
    run.font.size = Pt(18)

    for text in ["주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 기준",
                 "Web Application(웹) 21개 점검 항목",
                 f"점검일: {date_str}  |  대상: {target_url}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # === 1. 점검 개요 ===
    doc.add_heading("1. 점검 개요", level=1)
    add_kv_table([
        ("점검 대상", target_url),
        ("점검 도구", "OWASP ZAP + Ollama Gemma 4 (로컬 AI 분석)"),
        ("점검 방법", "ZAP Spider 크롤링 → Passive/Active Scan → 수동 점검 → AI 분석"),
        ("점검 기준", "2026 주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드 - Web Application(웹) 21개 항목"),
        ("소요 시간", duration),
    ])

    # === 2. 점검 결과 요약 ===
    doc.add_heading("2. 점검 결과 요약", level=1)

    v_counts = {"취약": 0, "주의": 0, "양호": 0, "수동점검 필요": 0}
    for f in findings:
        v_counts[f.verdict] = v_counts.get(f.verdict, 0) + 1

    # 판정 분포 테이블 (배경색 적용 — 핵심 요약이므로 필요)
    labels = ["취약", "주의", "양호", "수동점검 필요", "합계"]
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, label in enumerate(labels):
        hc = table.rows[0].cells[i]
        hc.text = label
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hc.paragraphs[0].runs[0].bold = True
        set_cell_shading(hc, "1A237E")
        hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    values = [v_counts["취약"], v_counts["주의"], v_counts["양호"],
              v_counts["수동점검 필요"], len(findings)]
    colors = ["FFEBEE", "FFF3E0", "E8F5E9", "F3E5F5", None]
    for i, (val, bg) in enumerate(zip(values, colors)):
        vc = table.rows[1].cells[i]
        vc.text = f"{val}건"
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc.paragraphs[0].runs[0].bold = True
        if bg:
            set_cell_shading(vc, bg)

    doc.add_paragraph()

    # 전체 항목 판정 테이블
    p = doc.add_paragraph()
    run = p.add_run("전체 항목별 판정 현황:")
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(findings), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(["코드", "점검 항목", "판정", "점검 방법"]):
        hc = table.rows[0].cells[i]
        hc.text = header
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hc.paragraphs[0].runs[0].bold = True
        set_cell_shading(hc, "1A237E")
        hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, f in enumerate(findings):
        row = table.rows[row_idx + 1]
        row.cells[0].text = f.code
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = f.name
        # 판정 셀 — 색상 텍스트 적용
        vc = row.cells[2]
        vc.text = f.verdict
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc.paragraphs[0].runs[0].bold = True
        vc.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(verdict_color(f.verdict))
        row.cells[3].text = f.scan_method_desc or ""

    doc.add_page_break()

    # === 3. 항목별 상세 결과 ===
    doc.add_heading("3. 점검 항목별 상세 결과", level=1)

    for f in findings:
        doc.add_heading(f"{f.code} - {f.full_name}", level=2)
        rows = [("항목 코드", f.code), ("판정", f.verdict),
                ("점검 방법", f.scan_method_desc), ("상세 결과", f.detail)]
        if f.verdict in ("취약", "주의"):
            rows.append(("조치 방안", f.remediation))
        add_kv_table(rows)

        doc.add_paragraph()

    doc.add_page_break()

    # === 4. ZAP 자동화 스캔 상세 ===
    doc.add_heading("4. ZAP 자동화 스캔 상세 결과", level=1)

    alerts_by_type = {}
    for f in findings:
        for a in f.zap_alerts[:5]:
            key = a.get("alert", "")
            if key not in alerts_by_type:
                alerts_by_type[key] = {**a, "urls": [], "count": 0}
            alerts_by_type[key]["urls"].append(a.get("url", "")[:80])
            alerts_by_type[key]["count"] += 1

    for a in alerts_by_type.values():
        doc.add_heading(f"{a['alert']} ({a.get('risk','')} - CWE-{a.get('cweid','')})", level=2)
        add_kv_table([
            ("위험도", a.get("risk", "")),
            ("CWE", f"CWE-{a.get('cweid', '')}"),
            ("탐지 건수", f"{a['count']}건"),
            ("영향 URL", ", ".join(a["urls"][:3])),
            ("설명", a.get("description", "")[:200]),
        ])
        doc.add_paragraph()

    doc.add_page_break()

    # === 5. 종합 의견 ===
    doc.add_heading("5. 종합 의견 및 권고사항", level=1)
    for line in summary.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("※ 본 보고서는 자동화 도구(OWASP ZAP)와 로컬 AI(Ollama Gemma 4)를 활용한 진단 결과이며, "
                     "정밀 진단을 위해 수동 점검이 추가로 필요합니다.")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)


# ============================================================
# 메인 실행
# ============================================================
def run_scan(target_url: str, skip_active=False, auth_config=None):
    start_time = time.time()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 출력 폴더 생성
    os.makedirs("logs", exist_ok=True)
    os.makedirs("docs", exist_ok=True)
    zap_log_dir = os.path.join("logs", "zap")
    os.makedirs(zap_log_dir, exist_ok=True)

    # 로그 파일을 타임스탬프별로 설정
    log_path = os.path.join("logs", f"scan_{timestamp}.log")
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.INFO)
    fh.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
    logger.addHandler(fh)

    logger.info("=" * 50)
    logger.info("스캔 시작 - 대상: %s", target_url)
    logger.info("=" * 50)
    print("=" * 64)
    print("  주요정보통신기반시설 웹 취약점 자동 진단 시스템 v2.0")
    print(f"  기준: 2026 가이드 - Web Application(웹) 21개 항목")
    print(f"  대상: {target_url}")
    print(f"  시작: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 64)

    # [1단계] 연결 확인
    stage_start = time.time()
    logger.info("[1단계] 연결 확인")
    print("\n[1단계] 연결 확인")
    zap = ZAPScanner()
    gemma = GemmaAnalyzer()
    if not zap.check():
        logger.error("[1단계] 실패 (%.1f초) - ZAP 연결 불가", time.time() - stage_start)
        print("\n  ZAP 실행 필요:")
        print("  docker run -u zap -p 8080:8080 zaproxy/zap-stable zap.sh -daemon -host 0.0.0.0 -port 8080 -config api.disablekey=true")
        return
    if not gemma.check():
        logger.error("[1단계] 실패 (%.1f초) - Ollama 연결 불가", time.time() - stage_start)
        print(f"\n  Ollama 실행 필요: ollama pull {OLLAMA_MODEL} && ollama serve")
        return
    logger.info("[1단계] 완료 (%.1f초) - ZAP, Ollama 연결 성공", time.time() - stage_start)

    context_name = f"WebScan-{timestamp}"
    context_id = ""

    # [3단계] 컨텍스트 & 인증 설정
    stage_start = time.time()
    logger.info("[3단계] 컨텍스트 설정")
    print("\n[3단계] 컨텍스트 설정")
    context_id = zap.create_context(context_name)
    regex = re.sub(r'(https?://)', r'\1', target_url).replace('.', '\\.') + ".*"
    zap.include_in_context(context_name, regex)

    if auth_config:
        logger.info("[3단계] 인증 설정")
        print("  [인증] 설정 중...")
        if auth_config.get("api_backend"):
            backend_regex = auth_config["api_backend"].replace('.', '\\.').replace(':', '\\:') + ".*"
            zap.include_in_context(context_name, backend_regex)
        zap.setup_auth(context_id, auth_config["login_url"], auth_config["login_data"],
                       auth_config.get("auth_type", "jsonBasedAuthentication"))
        zap.set_indicators(context_id,
                          auth_config.get("logged_in", "\\Qlogout\\E"),
                          auth_config.get("logged_out", "\\Qlogin\\E"))
        zap.create_user(context_id, auth_config["username"], auth_config["password"])

    auth_status = "설정" if auth_config else "미설정"
    logger.info("[3단계] 완료 (%.1f초) - 컨텍스트 ID:%s, 인증: %s", time.time() - stage_start, context_id, auth_status)

    # [4단계] Spider
    stage_start = time.time()
    logger.info("[4단계] Spider 크롤링")
    print("\n[4단계] Spider 크롤링")
    urls = zap.spider(target_url, context_name)
    logger.info("[4단계] 완료 (%.1f초) - %d개 URL 발견", time.time() - stage_start, len(urls))

    try:
        with open(os.path.join(zap_log_dir, f"spider_urls_{timestamp}.json"), "w", encoding="utf-8") as f:
            json.dump(urls, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("Spider URL 저장 실패: %s", e)

    # [5단계] Passive + Active Scan
    stage_start = time.time()
    logger.info("[5단계] Passive Scan + Active Scan")
    print("\n[5단계] Passive Scan + Active Scan")
    passive_ok = zap.wait_passive()
    if not skip_active:
        zap.active_scan(target_url, context_id=context_id)
        zap.wait_passive()
        active_status = "Active 완료"
    else:
        logger.info("[Active Scan] 건너뜀 (--skip-active)")
        print("  [Active Scan] 건너뜀 (--skip-active)")
        active_status = "Active 건너뜀"
    passive_status = "Passive 완료" if passive_ok else "Passive 타임아웃"
    logger.info("[5단계] 완료 (%.1f초) - %s, %s", time.time() - stage_start, passive_status, active_status)

    # [6단계] 결과 수집
    stage_start = time.time()
    logger.info("[6단계] ZAP 결과 수집")
    print("\n[6단계] ZAP 결과 수집")
    all_alerts = zap.get_alerts(target_url)
    logger.info("총 ZAP 경고: %d건", len(all_alerts))
    print(f"  총 ZAP 경고: {len(all_alerts)}건")
    summary = zap.get_alerts_summary(target_url)
    for k, v in (summary or {}).items():
        print(f"    {k}: {v}")
    logger.info("[6단계] 완료 (%.1f초) - %d건 경고 수집", time.time() - stage_start, len(all_alerts))

    try:
        with open(os.path.join(zap_log_dir, f"alerts_raw_{timestamp}.json"), "w", encoding="utf-8") as f:
            json.dump(all_alerts, f, ensure_ascii=False, indent=2)
        with open(os.path.join(zap_log_dir, f"alerts_summary_{timestamp}.json"), "w", encoding="utf-8") as f:
            json.dump(summary or {}, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("ZAP 경고 저장 실패: %s", e)

    # [7단계] 수동 점검
    stage_start = time.time()
    logger.info("[7단계] 수동 점검")
    print("\n[7단계] 수동 점검")
    manual_results = {}
    print("  [DI] 디렉터리 인덱싱 점검...")
    manual_results["DI"] = zap.manual_check_directories(target_url)
    logger.info("[7단계] DI 점검 완료 (%d건)", len(manual_results["DI"]))
    print("  [EP] 에러 페이지 점검...")
    manual_results["EP"] = zap.manual_check_error_pages(target_url)
    logger.info("[7단계] EP 점검 완료 (%d건)", len(manual_results["EP"]))
    print("  [AE] 관리자 페이지 노출 점검...")
    manual_results["AE"] = zap.manual_check_admin_pages(target_url)
    logger.info("[7단계] AE 점검 완료 (%d건)", len(manual_results["AE"]))
    print("  [WM] HTTP Method 점검...")
    manual_results["WM"] = zap.manual_check_http_methods(target_url)
    logger.info("[7단계] WM 점검 완료 (%d건)", len(manual_results["WM"]))
    print("  [SN] 보안 헤더 점검...")
    manual_results["SN"] = zap.manual_check_security_headers(target_url)
    logger.info("[7단계] SN 점검 완료")
    print("  [CC] 쿠키 보안 속성 점검...")
    manual_results["CC"] = zap.manual_check_cookies(target_url)
    logger.info("[7단계] CC 점검 완료")
    logger.info("[7단계] 완료 (%.1f초) - 6개 항목 수동 점검", time.time() - stage_start)

    try:
        with open(os.path.join(zap_log_dir, f"manual_checks_{timestamp}.json"), "w", encoding="utf-8") as f:
            json.dump(manual_results, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("수동 점검 결과 저장 실패: %s", e)
    logger.info("ZAP 결과 저장: %s", zap_log_dir)

    # 매핑
    alert_mapping = map_alerts_to_items(all_alerts)

    # [8단계] AI 분석 & 보고서
    stage_start = time.time()
    logger.info("[8단계] AI 분석 및 보고서 생성")
    print("\n[8단계] AI 분석 및 보고서 생성")
    findings = []
    for item in KISA_2026_ITEMS:
        mapped_alerts = alert_mapping.get(item["code"], [])
        mr = manual_results.get(item["code"])
        print(f"  [{item['code']}] {item['name']} (경고:{len(mapped_alerts)}건)...", end=" ")

        ai = gemma.analyze_item(item, mapped_alerts, mr)

        findings.append(Finding(
            code=item["code"],
            name=item["name"],
            full_name=item["full_name"],
            importance=item["importance"],
            verdict=ai.get("verdict", "수동점검 필요"),
            scan_method_desc=ai.get("scan_method_desc", ""),
            detail=ai.get("detail", ""),
            remediation=ai.get("remediation", ""),
            zap_alerts=mapped_alerts,
        ))
        logger.info("[%s] %s → %s (경고:%d건)", item['code'], item['name'], findings[-1].verdict, len(mapped_alerts))
        print(f"→ {findings[-1].verdict}")

    # 총괄 요약
    logger.info("종합 의견 생성 시작")
    print("\n  [종합 의견 생성 중...]")
    exec_summary = gemma.generate_summary(findings, target_url)
    if exec_summary.startswith("AI 분석 실패"):
        logger.error("종합 의견 생성 실패: %s", exec_summary)
    else:
        logger.info("종합 의견 생성 완료 (%d자)", len(exec_summary))

    v8 = sum(1 for f in findings if f.verdict == "취약")
    w8 = sum(1 for f in findings if f.verdict == "주의")
    s8 = sum(1 for f in findings if f.verdict == "양호")
    m8 = sum(1 for f in findings if f.verdict == "수동점검 필요")
    logger.info("[8단계] 완료 (%.1f초) - %d개 항목 분석 (취약:%d, 주의:%d, 양호:%d, 수동점검:%d), 종합 의견 %d자",
                time.time() - stage_start, len(findings), v8, w8, s8, m8, len(exec_summary))

    elapsed = time.time() - start_time
    scan_meta = {"duration": f"{int(elapsed//60)}분 {int(elapsed%60)}초"}

    # JSON 저장
    json_path = os.path.join("docs", f"vuln_report_{timestamp}.json")
    json_data = generate_json_report(findings, target_url, exec_summary, scan_meta)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)
    logger.info("JSON 보고서 저장: %s", json_path)

    # DOCX 보고서 생성
    docx_path = os.path.join("docs", f"vuln_report_{timestamp}.docx")
    print(f"\n  DOCX 생성 중...")
    try:
        generate_docx_report(findings, target_url, exec_summary, scan_meta,
                             os.path.abspath(docx_path))
        logger.info("DOCX 생성 완료: %s", docx_path)
        print(f"  DOCX 생성 완료: {docx_path}")
    except Exception as e:
        logger.error("DOCX 생성 실패: %s", e)
        print(f"  DOCX 생성 실패: {e}")

    # 결과 출력
    v = sum(1 for f in findings if f.verdict == "취약")
    w = sum(1 for f in findings if f.verdict == "주의")
    s = sum(1 for f in findings if f.verdict == "양호")
    m = sum(1 for f in findings if f.verdict == "수동점검 필요")

    logger.info("=" * 50)
    logger.info("진단 완료 - 취약: %d | 주의: %d | 양호: %d | 수동점검: %d", v, w, s, m)
    logger.info("소요시간: %s", scan_meta['duration'])
    logger.info("보고서: JSON=%s, DOCX=%s", json_path, docx_path)
    logger.info("ZAP 원시 데이터: %s", zap_log_dir)
    logger.info("=" * 50)
    print("\n" + "=" * 64)
    print("  진단 완료")
    print("=" * 64)
    print(f"  점검 항목: 21개 | 취약: {v} | 주의: {w} | 양호: {s} | 수동점검: {m}")
    print(f"  소요시간: {scan_meta['duration']}")
    print(f"  DOCX: {docx_path}")
    print(f"  JSON: {json_path}")
    print(f"  LOG:  {log_path}")
    print(f"  ZAP:  {zap_log_dir}")

    return docx_path, json_path


# ============================================================
# CLI
# ============================================================
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="주요정보통신기반시설 웹 취약점 자동 진단 v2.0 (2026 가이드 21개 항목)")
    parser.add_argument("target", help="진단 대상 URL")
    parser.add_argument("--skip-active", action="store_true")
    parser.add_argument("--zap-url", default=ZAP_API_URL)
    parser.add_argument("--zap-key", default=ZAP_API_KEY)
    parser.add_argument("--ollama-url", default=OLLAMA_URL)
    parser.add_argument("--model", default=OLLAMA_MODEL)
    # 인증 옵션
    parser.add_argument("--login-url", help="로그인 API URL")
    parser.add_argument("--login-data", help='로그인 JSON (예: {"email":"{%%username%%}","password":"{%%password%%}"})')
    parser.add_argument("--username", help="로그인 ID")
    parser.add_argument("--password", help="로그인 PW")
    parser.add_argument("--logged-in", default="\\Qlogout\\E|\\Q로그아웃\\E")
    parser.add_argument("--logged-out", default="\\Qlogin\\E|\\Q로그인\\E")
    parser.add_argument("--api-backend", help="API 백엔드 도메인 (프론트와 다를 때)")

    args = parser.parse_args()
    ZAP_API_URL = args.zap_url
    ZAP_API_KEY = args.zap_key
    OLLAMA_URL = args.ollama_url
    OLLAMA_MODEL = args.model

    auth = None
    if args.login_url and args.username:
        auth = {
            "login_url": args.login_url,
            "login_data": args.login_data or f'{{"email":"{{%username%}}","password":"{{%password%}}"}}',
            "username": args.username,
            "password": args.password,
            "logged_in": args.logged_in,
            "logged_out": args.logged_out,
            "api_backend": args.api_backend,
        }

    run_scan(args.target, skip_active=args.skip_active, auth_config=auth)
